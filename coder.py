import os
from docx import Document
from docx.shared import Pt

def clear_formatting_and_set_font(input_folder, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):
            input_filepath = os.path.join(input_folder, filename)
            output_filepath = os.path.join(output_folder, filename)

            try:
                document = Document(input_filepath)
                for paragraph in document.paragraphs[1:]:
                    paragraph.style = document.styles['Normal']
                    paragraph.style.font.name = 'Calibri'
                    paragraph.style.font.size = Pt(11)
                    paragraph.style.font.bold = False
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'

                if document.paragraphs:
                    first_paragraph = document.paragraphs[0]
                    first_paragraph.style.font.size = Pt(14)
                    first_paragraph.style.font.bold = True

                for paragraph in document.paragraphs:
                    if "?" in paragraph.text and paragraph.text.strip():
                        paragraph.text = f"<h3>{paragraph.text}</h3>"
                        paragraph.style.font.size = Pt(12)
                        paragraph.style.font.bold = True

                for paragraph in document.paragraphs:
                    if paragraph.text.strip() and len(paragraph.text.split()) < 7:
                        paragraph.text = f"<h2>{paragraph.text}</h2>"
                        paragraph.style.font.size = Pt(13)
                        paragraph.style.font.bold = True

                for paragraph in document.paragraphs:
                    if not paragraph.text.startswith("<h") and paragraph.text.strip():
                        paragraph.text = f"<p>{paragraph.text}</p>"
                        paragraph.style.font.bold = False

                document.save(output_filepath)
                print(f"Formatting cleared, font set, and paragraphs modified for {filename}.")

            except Exception as e:
                print(f"Error processing {filename}: {e}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("Usage: coder.py <input_folder> <output_folder>")
        sys.exit(1)

    input_folder = sys.argv[1]
    output_folder = sys.argv[2]
    clear_formatting_and_set_font(input_folder, output_folder)
